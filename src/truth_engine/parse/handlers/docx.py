"""DOCX via python-docx: paragraph text, heading structure, tables, and
document core properties (author/created/modified)."""

from __future__ import annotations

from pathlib import Path

from truth_engine.parse.registry import register
from truth_engine.parse.types import ParsedDocument, ParsedTable

PARSER_NAME = "python-docx"
PARSER_VERSION = "1"


@register("docx")
def parse_docx(path: Path) -> ParsedDocument:
    import docx  # lazy: pipeline extra

    document = docx.Document(str(path))

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    raw_text = "\n".join(paragraphs) if paragraphs else None

    headings = [
        p.text
        for p in document.paragraphs
        if p.text.strip() and p.style is not None and p.style.name.startswith("Heading")
    ]

    tables: list[ParsedTable] = []
    for table_num, table in enumerate(document.tables, start=1):
        rows_data = [[cell.text for cell in row.cells] for row in table.rows]
        if not rows_data:
            continue
        header, *rows = rows_data
        tables.append(
            ParsedTable(
                source=f"table_{table_num}",
                table_schema=header,
                rows=[dict(zip(header, row, strict=False)) for row in rows],
            )
        )

    core = document.core_properties
    embedded_metadata = {
        k: v
        for k, v in {
            "author": core.author,
            "title": core.title,
            "subject": core.subject,
            "last_modified_by": core.last_modified_by,
            "created": core.created.isoformat() if core.created else None,
            "modified": core.modified.isoformat() if core.modified else None,
        }.items()
        if v
    }

    return ParsedDocument(
        raw_text=raw_text,
        structure={"headings": headings} if headings else None,
        embedded_metadata=embedded_metadata or None,
        tables=tables,
        parser_name=PARSER_NAME,
        parser_version=PARSER_VERSION,
    )
